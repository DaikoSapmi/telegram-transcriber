"""Generate clean TXT and Word results from a local transcription."""

from __future__ import annotations

import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from config.settings import Settings
from src.transcription_types import TranscriptionResult


class DocumentGenerator:
    """Create deliverable files without fabricated speaker labels."""

    def __init__(self, output_dir: str | Path):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate(
        self,
        result: TranscriptionResult,
        original_filename: str,
        output_format: str,
    ) -> list[Path]:
        if output_format not in {"txt", "docx", "both"}:
            raise ValueError(f"Ugyldig resultatformat: {output_format}")
        paths = []
        if output_format in {"txt", "both"}:
            paths.append(self.generate_txt(result, original_filename))
        if output_format in {"docx", "both"}:
            paths.append(self.generate_docx(result, original_filename))
        return paths

    def generate_txt(self, result: TranscriptionResult, original_filename: str) -> Path:
        """Write plain UTF-8 transcript text, with no technical metadata."""
        path = (
            self.output_dir
            / f"{self._base_name(original_filename, result.language)}.txt"
        )
        text = result.text.strip()
        path.write_text(f"{text}\n" if text else "", encoding="utf-8")
        return path

    def generate_docx(
        self, result: TranscriptionResult, original_filename: str
    ) -> Path:
        document = Document()
        self._setup_document(document)

        title = document.add_heading("Transkripsjon", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        generated_at = datetime.now(timezone.utc).astimezone()
        run = subtitle.add_run(generated_at.strftime("%d.%m.%Y %H:%M"))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(110, 110, 110)

        document.add_heading("Informasjon", level=1)
        metadata = (
            ("Opprinnelig fil", Path(original_filename).name),
            ("Talespråk", Settings.get_language_name(result.language)),
            ("Modell", result.model_name),
            ("Varighet", self._format_timestamp(result.duration_seconds)),
        )
        for label, value in metadata:
            paragraph = document.add_paragraph()
            paragraph.add_run(f"{label}: ").bold = True
            paragraph.add_run(value)

        document.add_heading("Transkripsjon", level=1)
        for segment in result.segments:
            if not segment.text.strip():
                continue
            paragraph = document.add_paragraph()
            timestamp = paragraph.add_run(f"[{self._format_timestamp(segment.start)}] ")
            timestamp.font.name = "Courier New"
            timestamp.font.size = Pt(9)
            timestamp.font.color.rgb = RGBColor(110, 110, 110)
            paragraph.add_run(segment.text.strip())

        path = (
            self.output_dir
            / f"{self._base_name(original_filename, result.language)}.docx"
        )
        document.save(path)
        return path

    @staticmethod
    def _setup_document(document: Document) -> None:
        normal = document.styles["Normal"]
        normal.font.name = "Aptos"
        normal.font.size = Pt(11)
        for section in document.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.85)
            section.right_margin = Inches(0.85)

    @staticmethod
    def _format_timestamp(seconds: float) -> str:
        value = max(0, int(seconds))
        hours, remainder = divmod(value, 3600)
        minutes, secs = divmod(remainder, 60)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"

    @staticmethod
    def _base_name(original_filename: str, language: str) -> str:
        stem = Path(original_filename).stem.strip() or "lydopptak"
        stem = re.sub(r"[^\w.\-ÁČĐŊŠŽáčđŋšž]+", "_", stem, flags=re.UNICODE)
        language_name = (
            "nordsamisk"
            if language == "sme"
            else "norsk"
            if language == "no"
            else "automatisk"
        )
        today = datetime.now(timezone.utc).astimezone()
        return f"{stem}_{today:%Y-%m-%d}_{language_name}"
